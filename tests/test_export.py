from __future__ import annotations

import csv
import json
import base64
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

from prospectus_sentence_indexer.export import Exporter
from prospectus_sentence_indexer.models import BlockType, QAReport, Sentence, Summary


def _sentence(text: str = "Sentence.") -> Sentence:
    return Sentence(
        sentence_id="S000001",
        block_id="B000001",
        part="document.xml",
        block_type=BlockType.PARAGRAPH,
        heading_level=1,
        heading_path="1 Business",
        text=text,
        qa_flags=["TOO_SHORT"],
    )


def test_export_csv_format(tmp_path: Path) -> None:
    exporter = Exporter(str(tmp_path))
    output = exporter.export_csv([_sentence()])

    with Path(output).open(encoding="utf-8-sig") as f:
        rows = list(csv.reader(f))

    assert rows[0][0] == "sentence_id"
    assert rows[1][0] == "S000001"
    assert rows[1][6] == "Sentence."


def test_export_docx_table(tmp_path: Path) -> None:
    exporter = Exporter(str(tmp_path))
    output = exporter.export_docx([_sentence()], input_name="Project Star-IO.pdf", timestamp=datetime(2026, 3, 6, 10, 12, 13))

    doc = Document(output)
    table = doc.tables[0]
    assert Path(output).name == "Project_Star-IO_20260306_101213.docx"
    assert table.rows[0].cells[0].text == "sentence_id"
    assert table.rows[0].cells[1].text == "text+image"
    assert table.rows[0].cells[2].text == "qa_flags"
    assert len(table.columns) == 3
    assert table.rows[1].cells[0].text == "S000001"
    assert doc.sections[0].orientation == WD_ORIENT.LANDSCAPE
    assert doc.sections[0].page_width.inches == 11
    assert doc.sections[0].page_height.inches == 8.5
    assert doc.styles["Normal"].font.name == "Times New Roman"
    for cell in table.rows[1].cells:
        tc_borders = cell._tc.tcPr.find(qn("w:tcBorders"))
        assert tc_borders is not None
        assert tc_borders.find(qn("w:top")).get(qn("w:sz")) == "4"


def test_export_docx_heading_merges_three_cells(tmp_path: Path) -> None:
    exporter = Exporter(str(tmp_path))
    heading = Sentence(
        sentence_id="",
        block_id="B000000",
        part="document.xml",
        block_type=BlockType.HEADING,
        heading_level=1,
        heading_path="Heading",
        text="INDUSTRY OVERVIEW",
        qa_flags=[],
    )
    output = exporter.export_docx(
        [heading, _sentence()],
        input_name="sample.pdf",
        timestamp=datetime(2026, 3, 6, 10, 12, 13),
    )

    doc = Document(output)
    table = doc.tables[0]
    heading_row = table.rows[1]
    assert heading_row.cells[0].text == "INDUSTRY OVERVIEW"
    grid_span = heading_row.cells[0]._tc.tcPr.find(qn("w:gridSpan"))
    assert grid_span is not None
    assert grid_span.get(qn("w:val")) == "3"


def test_export_docx_image_embeds_in_text_cell_and_hides_placeholder_name(tmp_path: Path) -> None:
    image_path = tmp_path / "sample.png"
    image_path.write_bytes(base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+XGZ0AAAAASUVORK5CYII="
    ))
    sentence = Sentence(
        sentence_id="S000002",
        block_id="B000002",
        part="document.xml",
        block_type=BlockType.IMAGE,
        heading_level=0,
        heading_path="",
        text="[IMAGE] images/sample.png",
        image_path=str(image_path),
        qa_flags=[],
    )

    exporter = Exporter(str(tmp_path))
    output = exporter.export_docx([sentence])

    doc = Document(output)
    table = doc.tables[0]
    assert table.rows[1].cells[1].text == ""
    assert len(doc.inline_shapes) == 1


def test_export_qa_report(tmp_path: Path) -> None:
    exporter = Exporter(str(tmp_path))
    output = exporter.export_qa_report(
        [
            QAReport(
                block_id="B000001",
                sentence_id="S000001",
                part="document.xml",
                block_type=BlockType.PARAGRAPH,
                heading_path="1 Business",
                issue="TOO_SHORT",
                details="short",
                raw_excerpt="abc",
            )
        ]
    )

    with Path(output).open(encoding="utf-8-sig") as f:
        rows = list(csv.reader(f))

    assert rows[0][0] == "block_id"
    assert rows[1][5] == "TOO_SHORT"


def test_export_summary_json(tmp_path: Path) -> None:
    exporter = Exporter(str(tmp_path))
    output = exporter.export_summary_json(
        Summary(
            input_file="sample.docx",
            processing_time_seconds=1.2,
            total_blocks=10,
            total_sentences=12,
            block_type_counts={"paragraph": 10},
            qa_issue_counts={"TOO_SHORT": 2},
            reconstruct_fail_count=1,
            fallback_sentence_count=1,
        )
    )

    data = json.loads(Path(output).read_text(encoding="utf-8"))
    assert data["input_file"] == "sample.docx"
    assert data["qa_issue_counts"]["TOO_SHORT"] == 2


def test_export_special_chars(tmp_path: Path) -> None:
    exporter = Exporter(str(tmp_path))
    special = 'Quote " and comma, and unicode 測試'
    output = exporter.export_csv([_sentence(text=special)])

    with Path(output).open(encoding="utf-8-sig") as f:
        rows = list(csv.reader(f))

    assert rows[1][6] == special


def test_export_csv_includes_image_path(tmp_path: Path) -> None:
    exporter = Exporter(str(tmp_path))
    output = exporter.export_csv([_sentence(text="Image row.")])

    with Path(output).open(encoding="utf-8-sig") as f:
        rows = list(csv.DictReader(f))

    assert "image_path" in rows[0]
