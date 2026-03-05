#!/usr/bin/env python3
"""
Convert DOCX to Markdown and JSON
"""

from docx import Document
import json
import os
import sys


def convert_docx_to_markdown_and_json(docx_path, output_dir):
    """Convert DOCX file to Markdown and JSON formats"""
    
    # 读取 docx 文件
    doc = Document(docx_path)
    
    # 创建输出目录
    os.makedirs(output_dir, exist_ok=True)
    os.makedirs(f'{output_dir}/images', exist_ok=True)
    
    # 提取文本内容
    content_list = []
    markdown_lines = []
    
    for para_idx, para in enumerate(doc.paragraphs):
        if para.text.strip():
            # 检测标题级别
            level = 0
            if para.style.name.startswith('Heading'):
                try:
                    level = int(para.style.name.replace('Heading ', ''))
                except:
                    pass
            
            # 添加到 content_list
            content_list.append({
                'type': 'text',
                'text': para.text,
                'text_level': level,
                'para_idx': para_idx
            })
            
            # 生成 Markdown
            if level > 0:
                heading_prefix = '#' * level
                markdown_lines.append(f'{heading_prefix} {para.text}')
            else:
                markdown_lines.append(para.text)
    
    # 提取表格
    for table_idx, table in enumerate(doc.tables):
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        
        content_list.append({
            'type': 'table',
            'table_idx': table_idx,
            'data': table_data
        })
        
        # 生成 Markdown 表格
        if table_data:
            markdown_lines.append('')  # 空行
            # 表头
            header = '| ' + ' | '.join(table_data[0]) + ' |'
            markdown_lines.append(header)
            # 分隔符
            separator = '| ' + ' | '.join(['---'] * len(table_data[0])) + ' |'
            markdown_lines.append(separator)
            # 数据行
            for row in table_data[1:]:
                row_str = '| ' + ' | '.join(row) + ' |'
                markdown_lines.append(row_str)
            markdown_lines.append('')  # 空行
    
    # 保存 JSON
    json_path = f'{output_dir}/content_list.json'
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(content_list, f, ensure_ascii=False, indent=2)
    
    # 保存 Markdown
    md_path = f'{output_dir}/document.md'
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(markdown_lines))
    
    # 统计信息
    text_count = len([c for c in content_list if c['type'] == 'text'])
    table_count = len([c for c in content_list if c['type'] == 'table'])
    
    print(f'转换完成！')
    print(f'JSON 文件: {json_path}')
    print(f'Markdown 文件: {md_path}')
    print(f'段落数: {text_count}')
    print(f'表格数: {table_count}')
    
    return {
        'json_path': json_path,
        'md_path': md_path,
        'text_count': text_count,
        'table_count': table_count
    }


if __name__ == '__main__':
    docx_path = '/Users/ethan/Documents/trae_projects/HK_IPO_docx/docx/Project_Star-IO_0305.docx'
    output_dir = '/Users/ethan/Documents/trae_projects/HK_IPO_docx/output/Project_Star-IO_0305'
    
    result = convert_docx_to_markdown_and_json(docx_path, output_dir)
