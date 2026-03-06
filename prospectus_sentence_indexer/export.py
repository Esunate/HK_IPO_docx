from __future__ import annotations

import csv
import json
from dataclasses import asdict
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.shared import Pt

from .models import BlockType, QAReport, Sentence, Summary


class Exporter:
    def __init__(self, output_dir: str):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def export_csv(self, sentences: list[Sentence]) -> str:
        output = self.output_dir / "sentences.csv"
        with output.open("w", newline="", encoding="utf-8-sig") as f:
            writer = csv.writer(f)
            writer.writerow([
                "sentence_id",
                "block_id",
                "part",
                "block_type",
                "heading_level",
                "heading_path",
                "text",
                "image_path",
                "qa_flags",
            ])
            for sentence in sentences:
                writer.writerow([
                    sentence.sentence_id,
                    sentence.block_id,
                    sentence.part,
                    sentence.block_type.value,
                    sentence.heading_level,
                    sentence.heading_path,
                    sentence.text,
                    sentence.image_path or "",
                    "|".join(sentence.qa_flags),
                ])
        return str(output)

    def export_docx(
        self,
        sentences: list[Sentence],
        input_name: str = "sentences",
        timestamp: datetime | None = None,
    ) -> str:
        output = self.output_dir / self._docx_filename(input_name, timestamp)
        document = Document()
        self._set_default_font(document)
        self._set_landscape_letter(document)
        table = document.add_table(rows=1, cols=3)
        self._set_table_borders(table)
        headers = ["sentence_id", "text+image", "qa_flags"]
        for idx, header in enumerate(headers):
            table.rows[0].cells[idx].text = header
            self._set_cell_borders(table.rows[0].cells[idx])

        for sentence in sentences:
            row = table.add_row()
            if sentence.block_type == BlockType.HEADING:
                merged_cell = row.cells[0].merge(row.cells[2])
                merged_cell.text = sentence.text.strip()
                self._set_cell_borders(merged_cell)
                continue

            cells = row.cells
            cells[0].text = sentence.sentence_id
            self._set_cell_borders(cells[0])
            self._write_text_image_cell(cells[1], sentence.text, sentence.image_path)
            self._set_cell_borders(cells[1])
            cells[2].text = "|".join(sentence.qa_flags)
            self._set_cell_borders(cells[2])

        document.save(output)
        return str(output)

    @staticmethod
    def _docx_filename(input_name: str, timestamp: datetime | None) -> str:
        safe_name = "_".join(Path(input_name).stem.split()) or "sentences"
        dt = timestamp or datetime.now()
        return f"{safe_name}_{dt.strftime('%Y%m%d_%H%M%S')}.docx"

    @staticmethod
    def _set_landscape_letter(document: Document) -> None:
        section = document.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)

    @staticmethod
    def _set_default_font(document: Document) -> None:
        style = document.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(10)
        r_pr = style.element.get_or_add_rPr()
        r_fonts = r_pr.rFonts
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.append(r_fonts)
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            r_fonts.set(qn(f"w:{attr}"), "Times New Roman")

    @staticmethod
    def _set_table_borders(table) -> None:
        table_pr = table._tbl.tblPr
        borders = table_pr.first_child_found_in("w:tblBorders")
        if borders is None:
            borders = OxmlElement("w:tblBorders")
            table_pr.append(borders)
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = borders.find(qn(f"w:{edge}"))
            if border is None:
                border = OxmlElement(f"w:{edge}")
                borders.append(border)
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "4")
            border.set(qn("w:color"), "000000")

    @staticmethod
    def _set_cell_borders(cell) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = tc_pr.first_child_found_in("w:tcBorders")
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            tc_pr.append(borders)
        for edge in ("top", "left", "bottom", "right"):
            border = borders.find(qn(f"w:{edge}"))
            if border is None:
                border = OxmlElement(f"w:{edge}")
                borders.append(border)
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "4")
            border.set(qn("w:color"), "000000")

    @staticmethod
    def _write_text_image_cell(cell, text: str, image_path: str | None) -> None:
        display_text = Exporter._display_text(text, image_path)
        if display_text:
            cell.paragraphs[0].text = display_text
        else:
            cell.paragraphs[0].text = ""

        if not image_path:
            return

        path = Path(image_path)
        if not path.exists():
            return

        paragraph = cell.add_paragraph() if display_text else cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(str(path), width=Inches(3.8))

    @staticmethod
    def _display_text(text: str, image_path: str | None) -> str:
        normalized = text.strip()
        if image_path and normalized.startswith("[IMAGE]"):
            return ""
        return normalized

    def export_qa_report(self, qa_reports: list[QAReport]) -> str:
        output = self.output_dir / "qa_report.csv"
        with output.open("w", newline="", encoding="utf-8-sig") as f:
            writer = csv.writer(f)
            writer.writerow([
                "block_id",
                "sentence_id",
                "part",
                "block_type",
                "heading_path",
                "issue",
                "details",
                "raw_excerpt",
            ])
            for report in qa_reports:
                writer.writerow([
                    report.block_id,
                    report.sentence_id or "",
                    report.part,
                    report.block_type.value,
                    report.heading_path,
                    report.issue,
                    report.details,
                    report.raw_excerpt,
                ])
        return str(output)

    def export_summary_json(self, summary: Summary) -> str:
        output = self.output_dir / "summary.json"
        with output.open("w", encoding="utf-8") as f:
            json.dump(asdict(summary), f, ensure_ascii=False, indent=2)
        return str(output)
